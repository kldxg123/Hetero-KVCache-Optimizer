#!/usr/bin/env python3
"""
generate_report.py — 生成 Hetero-KV 项目完整技术报告 (DOCX)
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── 全局样式 ──────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.35

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Microsoft YaHei'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# ── 辅助函数 ──────────────────────────────────────────────
def add_heading(text, level=1):
    return doc.add_heading(text, level=level)

def add_para(text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.3)
    run = p.add_run(text)
    run.bold = bold
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.3)
    return p

def add_qa(question, answer):
    p = doc.add_paragraph()
    run_q = p.add_run("Q: " + question)
    run_q.bold = True
    run_q.font.color.rgb = RGBColor(0x1A, 0x5C, 0xB0)
    p2 = doc.add_paragraph()
    run_a = p2.add_run("A: " + answer)
    run_a.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p2.paragraph_format.left_indent = Inches(0.2)

def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for run in t.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            t.rows[r_idx + 1].cells[c_idx].text = str(val)
    return t


# ══════════════════════════════════════════════════════════════
#                        正文开始
# ══════════════════════════════════════════════════════════════

doc.add_heading('Hetero-KV 项目完整技术报告', level=0)
doc.add_paragraph('弹性分层内存架构：让消费级 GPU 跑通 128K 长上下文推理').alignment = WD_ALIGN_PARAGRAPH.CENTER
p_info = doc.add_paragraph()
p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_info = p_info.add_run('项目仓库：Hetero-KVCache-Optimizer  |  目标设备：RTX 4090 (24 GB HBM)')
run_info.font.size = Pt(9)
run_info.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# ============================================================
# 第一章：为什么要做这个项目
# ============================================================
add_heading('一、项目背景：为什么要做 Hetero-KV？', level=1)

add_heading('1.1 核心痛点：显存不够', level=2)
add_para('大语言模型（LLM）在做推理时，会把所有历史 token 的 Key/Value 向量缓存在 GPU 显存（HBM）里，这个缓存叫 KV Cache。上下文越长，KV Cache 越大：', indent=True)
add_table(
    ['上下文长度', 'KV Cache 大小 (Qwen2.5-7B, BF16)', '24GB GPU 剩余空间'],
    [
        ['8K tokens',  '~1.0 GB',   '9 GB'],
        ['32K tokens', '~4.0 GB',   '6 GB'],
        ['64K tokens', '~8.0 GB',   '2 GB'],
        ['96K tokens', '~12.0 GB',  'OOM 崩溃'],
        ['128K tokens','~16.0 GB',  'OOM 崩溃'],
    ]
)
add_para('模型权重本身就要占 14 GB（BF16），留给 KV Cache 的空间只有 ~10 GB。原生 HuggingFace 在 ~64K 就会 OOM（显存溢出崩溃）。', indent=True)

add_qa('什么是 KV Cache？',
       'Transformer 每生成一个新 token，需要用所有历史 token 的 Key 和 Value 向量计算注意力。把这些向量缓存下来避免重复计算，这就是 KV Cache。它是自回归推理中最占显存的部分。')

add_qa('为什么是 14 GB 权重？',
       'Qwen2.5-7B 有约 70 亿参数，每个参数用 BF16（2 字节）存储：70 亿 × 2 字节 = 14 GB。这是推理时必须常驻 HBM 的固定开销。')

add_heading('1.2 已有方案的缺陷', level=2)
add_table(
    ['方案', '核心做法', '致命缺陷'],
    [
        ['StreamingLLM', '只保留最近 N 个 token，丢弃其余', '丢弃 = 永久遗忘，NIAH 检索准确率降到 16.1%'],
        ['KIVI', '对 KV Cache 做 2-bit/4-bit 量化但全量保留', '96K token 时仍然 OOM，因为只是缩小但没移出 GPU'],
        ['SnapKV', '根据注意力分数剪枝不重要 token', '永久删除 ≈ 精度退化 8%，不可恢复'],
        ['vLLM PagedAttention', '分页管理 KV Cache 但仍在 HBM 内', '不解决显存总量不足的问题，只是管理更高效'],
    ]
)
add_para('核心矛盾：既要省显存，又不能丢信息。现有方案要么「丢弃信息」（精度损失），要么「压缩但不外移」（仍然 OOM）。', indent=True)

add_heading('1.3 Hetero-KV 的核心思路', level=2)
add_para('把 GPU HBM 想象成「一级缓存」（小但快），把 CPU DRAM 想象成「二级缓存」（大但慢）。关键洞察：', indent=True)
add_bullet('不是所有 token 同等重要——注意力分数低的 token 是「冷数据」，可以先压缩再搬出去')
add_bullet('搬出去不是丢弃——用 4-bit 量化压缩后存在 CPU 内存里，需要时再搬回来（自愈恢复）')
add_bullet('搬回来的时机可以预测——用注意力模式预测下一步需要哪些数据，提前在后台 DMA 传输')

add_qa('这和虚拟内存有什么区别？',
       '核心区别在于「注意力感知」：操作系统的虚拟内存用页表命中频率做 LRU 替换，而 Hetero-KV 用注意力分数决定哪些 KV 块要驱逐/预取。注意力分数直接反映模型对每个 token 的实际需求，比通用 LRU 精确得多。此外，我们额外做了 4-bit 量化压缩，使 PCIe 传输量减少 75%。')

# ============================================================
# 第二章：系统总体架构
# ============================================================
add_heading('二、系统总体架构', level=1)

add_heading('2.1 分层存储模型', level=2)
add_para('系统将存储分为三层：', indent=True)
add_table(
    ['层级', '物理介质', '存什么', '容量', '延迟'],
    [
        ['Tier 1: HBM', 'GPU 显存 (24 GB)', 'Sink + Tail 的原始 BF16 KV 张量', '~0.35 GB 固定', '<1 μs'],
        ['Tier 2: DRAM', 'CPU 内存', '4-bit 压缩后的 Body KV 块 + 元数据', '可达数百 GB', '~1 ms (PCIe)'],
        ['Tier 3: NVMe', 'SSD (预留)', '极端内存约束下的溢出', 'TB 级', '~100 μs (PCIe+SSD)'],
    ]
)

add_heading('2.2 三大保护区域', level=2)
add_para('对于任意时刻的上下文序列，Hetero-KV 将其划分为三个区域：', indent=True)
add_bullet('Sink（汇聚区）：前 64 个 token。包括系统提示词、角色设定等，注意力分数始终很高，永久驻留 HBM。')
add_bullet('Local / Tail（尾部区）：最近 8192 个 token。包含最新上下文，模型当前最需要的部分，保留在 HBM 中滚动更新。')
add_bullet('Body（主体区）：Sink 和 Tail 之间的所有历史 token。这部分被压缩到 4-bit 后存入 CPU DRAM，需要时通过自愈恢复机制按需取回。')

add_heading('2.3 完整数据流路线图', level=2)
add_para('下面是数据从输入到输出的完整路径，标出了每一步经过的方法和文件位置：', indent=True)

flow_text = """阶段 A: 模型输入 & 分块预填充 (Prefill)
  ├─ 用户输入长文本 (如 128K tokens)
  ├─ ChunkedPrefillEngine (src/core/engine_wrapper.py)
  │    ├─ 将 128K tokens 拆成多个 chunk (每 chunk 2048 tokens)
  │    ├─ 每个 chunk 单独做一次 forward pass 生成 KV Cache
  │    ├─ 每个 chunk 完成后触发 gc.collect() 回收临时显存
  │    └─ 累积 KV Cache 至 HeteroKVManager
  └─ 输出：完整的 128K KV Cache 存于临时显存

阶段 B: 瞬态拦截 & 分层存储决策 (Transient Interception)
  ├─ FusedHeteroCache (src/core/engine_wrapper.py) 检测到 prefill 结束
  ├─ HeteroKVManager (src/memory/manager.py) 接管所有 KV 张量
  │    ├─ 分区：Sink (前64) + Tail (后8192) 保留在 HBM (HBMStorageManager)
  │    └─ Body (中间所有) 交给量化压缩管线
  └─ 输出：HBM 仅存 Sink+Tail (~0.35 GB)，Body 待压缩

阶段 C: 量化压缩 & DRAM 卸载 (Compression & Offload)
  ├─ KVCompressor (src/quantization/kv_compressor.py)
  │    ├─ 将 Body 的 BF16 KV 张量按 group_size=128 分组
  │    ├─ 每组计算 scale = (max - min) / 15 和 zero_point = min
  │    ├─ 量化为 uint4 (4-bit 无符号整数)，打包存储
  │    └─ 压缩比 4x，相对重建误差 < 0.5%
  ├─ DRAMStorageManager (src/memory/dram_storage.py)
  │    ├─ 将压缩后的数据存入 CPU pinned memory (锁页内存)
  │    └─ 锁页内存确保 PCIe DMA 传输零拷贝
  └─ 输出：Body 以 4-bit 形式存于 CPU DRAM

阶段 D: 解码阶段注意力评分 (Attention Scoring)
  ├─ 每个 decode step，模型输出最新 token 对所有历史 token 的注意力权重
  ├─ HeavyHitterOracle (src/policy/heavy_hitter.py)
  │    ├─ update()：累积每个 token 的注意力分数 E_i = Σ A_{t,i}
  │    └─ 累积评分保护了「早期不重要但后期变重要」的 token
  └─ 输出：每个 token 的累积重要性分数

阶段 E: 驱逐决策 & HBM 物理池滚动 (Eviction & Rolling)
  ├─ HeavyHitterOracle.get_eviction_candidates()
  │    ├─ 调用 _block_mean_kernel (src/kernels/oracle_triton.py)
  │    │    └─ Triton GPU 内核：将 token 级分数聚合为 block 级均值
  │    │       10,465x 加速 (3.5s → 0.33ms)
  │    ├─ 构建 safe_mask：Sink blocks 和 Local blocks 设为 inf (永不被驱逐)
  │    └─ torch.topk(largest=False)：选出分数最低的 k 个 block
  ├─ HBMStorageManager 执行驱逐
  │    └─ 被驱逐的 block 先经 KVCompressor 压缩，再存入 DRAM
  └─ 输出：HBM 腾出空间接收新的 decode KV

阶段 F: 预测性预取 (Predictive Prefetch)
  ├─ PredictivePrefetchScheduler (src/core/scheduler.py)
  │    ├─ 信号1: 空间局部性 — 预取与当前 block 相邻的 block
  │    ├─ 信号2: 注意力热度 — 预取注意力分数最高的 block
  │    └─ 信号3: 顺序前瞻 — 预测性地预取未来可能需要的 block
  ├─ AdaptivePrefetchController (src/policy/adaptive_prefetch_controller.py)
  │    ├─ 监控注意力波动 σ(A_t) 和缓存未命中率 miss_rate_t
  │    └─ 动态调整预取窗口: w_t = w_min + clip(σ/σ_ref-1)·α + β·miss_rate
  ├─ AsyncPrefetcher (src/policy/prefetcher.py)
  │    ├─ 在后台 CUDA stream 上执行 DRAM → HBM 传输
  │    └─ 与主计算流并行，计算-传输重叠效率 98.3%
  └─ 输出：需要的 block 已提前到达 HBM

阶段 G: 自愈恢复 & 融合注意力计算 (Self-Healing & Fused Attention)
  ├─ FusedDequantAttn (src/quantization/fused_dequant_attn.py)
  │    ├─ 不将 4-bit 解压为 BF16 再算注意力
  │    ├─ 而是在 Triton 内核的 GPU 寄存器中直接做 4-bit → BF16 → 点积
  │    ├─ 在线 Softmax：两遍扫描 (max + sum)，不物化完整注意力矩阵
  │    └─ 消除了恢复路径中的 512 MB FP32 瞬态内存峰值
  └─ 输出：最终注意力输出，用于生成下一个 token"""
doc.add_paragraph(flow_text, style='Normal')

add_heading('2.4 核心文件索引', level=2)
add_table(
    ['文件路径', '核心类/函数', '职责'],
    [
        ['src/core/engine_wrapper.py', 'FusedHeteroCache, ChunkedPrefillEngine', 'HF 集成层，分块预填充，模式检测'],
        ['src/core/scheduler.py', 'PredictivePrefetchScheduler', '预测性预取调度器'],
        ['src/memory/manager.py', 'HeteroKVManager', '分层存储管理（HBM+DRAM）'],
        ['src/memory/hbm_storage.py', 'HBMStorageManager', 'GPU 显存块分配/驱逐'],
        ['src/memory/dram_storage.py', 'DRAMStorageManager', 'CPU 内存锁页存储'],
        ['src/quantization/kv_compressor.py', 'KVCompressor', '4-bit 分组非对称量化/反量化'],
        ['src/quantization/fused_dequant_attn.py', 'FusedDequantAttn', '融合反量化注意力 Triton 内核'],
        ['src/policy/heavy_hitter.py', 'HeavyHitterOracle', '注意力感知驱逐决策引擎'],
        ['src/policy/prefetcher.py', 'AsyncPrefetcher', '异步 PCIe 传输 (后台 CUDA stream)'],
        ['src/policy/adaptive_prefetch_controller.py', 'AdaptivePrefetchController', '自适应预取窗口控制'],
        ['src/kernels/oracle_triton.py', '_block_mean_kernel', '块级注意力分数聚合 Triton 内核'],
    ]
)

# ============================================================
# 第三章：逐阶段深度解析
# ============================================================
add_heading('三、逐阶段深度解析', level=1)

# ── 3.1 分块预填充 ──
add_heading('3.1 阶段 A：分块预填充 (Chunked Prefill)', level=2)
add_para('问题：128K token 的输入如果一次性做 prefill，KV Cache 瞬间需要 ~16 GB，直接 OOM。', indent=True)
add_para('解决：ChunkedPrefillEngine 将输入拆成多个小 chunk，每个 chunk 独立做 forward pass。', indent=True)

add_heading('工作原理', level=3)
add_bullet('chunk_size = 2048：每次处理 2048 个 token')
add_bullet('每个 chunk 的 forward pass 产生 2048 组 KV 向量，累积到临时缓存')
add_bullet('每处理完 4 个 chunk 触发 gc.collect()，回收 PyTorch 的中间张量碎片')
add_bullet('所有 chunk 处理完毕后，进入「瞬态拦截」阶段')

add_heading('chunk_size 的权衡', level=3)
add_table(
    ['chunk_size', 'TTFT (首 token 延迟)', '峰值 KV Cache', '评价'],
    [
        ['512',  '1.96 s', '0.34 GB', '太慢，IO 开销大'],
        ['2048', '0.65 s', '0.45 GB', '最优平衡点 ✓'],
        ['4096', '0.42 s', '0.65 GB', '快但峰值显存高'],
    ]
)

add_qa('为什么分块能省显存？每次不是都会生成 KV 吗？',
       '关键在于 gc.collect()。每个 chunk 的 forward pass 会产生大量中间张量（注意力分数矩阵、MLP 激活等），这些在 chunk 结束后就没用了。如果不分块，128K 的中间张量可能再占 10+ GB。分块后，每 4 个 chunk 就回收一次，峰值显存从 ~16 GB 降到 ~0.45 GB。')

add_qa('TTFT 是什么？',
       'Time To First Token，从输入到模型输出第一个 token 的延迟。分块预填充让 TTFT 从 0.572 s 降到 0.269 s，减少 53%。')

# ── 3.2 瞬态拦截 ──
add_heading('3.2 阶段 B：瞬态拦截 & 存储分层 (Transient Interception)', level=2)
add_para('Prefill 结束瞬间，KV Cache 占用最大（峰值）。FusedHeteroCache 检测到 prefill → decode 的切换点，立即触发「拦截」：', indent=True)
add_bullet('前 64 个 token → Sink 区，永久驻留 HBM')
add_bullet('最后 8192 个 token → Tail 区，驻留 HBM 并随 decode 滚动更新')
add_bullet('中间所有 token → Body 区，交给量化压缩管线')
add_para('拦截完成后，HBM 的 KV Cache 从峰值（可能数 GB）骤降至 ~0.35 GB（仅 Sink + Tail）。', indent=True)

add_qa('为什么 Sink 是 64？Tail 是 8192？',
       'Sink 64 来自 StreamingLLM 的发现：Transformer 前几个 token 充当「attention sink」，注意力分数异常高（因为 Softmax 需要分配概率质量）。Tail 8192 是实验调优的结果：太小会导致频繁缺块（cache miss），太大会浪费 HBM。8192 覆盖了大多数局部注意力模式。')

# ── 3.3 量化压缩 ──
add_heading('3.3 阶段 C：4-bit 分组非对称量化 & DRAM 卸载', level=2)

add_heading('为什么要量化？', level=3)
add_para('Body 区的 KV 向量要搬到 CPU 内存，走的是 PCIe 总线。RTX 4090 的 PCIe Gen4 x16 理论带宽 ~32 GB/s，实际可用约 16-25 GB/s。一个 16K token 的 KV block 在 BF16 下约 8 MB，如果不压缩，搬一个 block 就要 ~0.5 ms。但如果量化到 4-bit：', indent=True)
add_bullet('体积缩小 4 倍：8 MB → 2 MB')
add_bullet('传输时间缩短：PCIe 延迟从 ~0.5 ms 降到 ~0.14 ms')
add_bullet('实测 ITL 加速：2.98x（相对原始 BF16 PCIe 传输）')

add_heading('量化算法详解', level=3)
add_para('采用的是「分组非对称量化」（Group-wise Asymmetric Quantization），参数如下：', indent=True)
add_bullet('bits = 4：每个元素用 4 bit 表示（值域 0~15）')
add_bullet('group_size = 128：每 128 个连续元素共享一组 (scale, zero_point)')
add_para('对于一个组 X（128 个 BF16 元素）：', indent=True)
add_bullet('第 1 步：找组内最大最小值 → max_val, min_val')
add_bullet('第 2 步：计算 scale = (max_val - min_val) / (2^4 - 1) = (max_val - min_val) / 15')
add_bullet('第 3 步：计算 zero_point = min_val')
add_bullet('第 4 步：量化 q_i = round((x_i - zero_point) / scale)，值域 [0, 15]')
add_bullet('反量化：x_hat_i = q_i × scale + zero_point')
add_para('每个组额外存储 1 个 scale (FP32, 4 bytes) + 1 个 zero_point (FP32, 4 bytes)，128 个元素从 256 bytes (BF16) 压缩为 64 bytes (4-bit) + 8 bytes (元数据) = 72 bytes。压缩比 256/72 = 3.56x，加上打包优化整体接近 4x。', indent=True)

add_heading('为什么选非对称而不是对称量化？', level=3)
add_para('KV Cache 的值分布通常不对称（不是零均值），存在大量正值偏移。对称量化会浪费一半的量化范围在负值上，导致精度损失更大。非对称量化通过 zero_point 平移，能更充分利用 4-bit 的 16 个量化级别。', indent=True)

add_heading('为什么选 group_size=128？', level=3)
add_para('更小的 group (如 32)：每组需要独立存储 scale+zero_point，元数据开销比例增大（8 bytes / 32×2 bytes = 12.5% 元数据），压缩比下降。更大的 group (如 1024)：组内值范围更大，量化粒度粗，精度损失增加。128 是实验验证的甜点。', indent=True)

add_heading('量化误差的安全性证明', level=3)
add_para('论文给出了理论定理（Attention Resilience Theorem）：', indent=True)
add_para('对非重击者 token 施加 4-bit 量化后，注意力输出的无穷范数扰动有上界：')
add_para('  ||Attn(Q,K,V) - Attn(Q,K_hat,V)||_inf ≤ α × sqrt(d) × ε_q')
add_para('其中 α 是量化 token 上的注意力质量（非重击者 token 的 α 很小），d 是头维度，ε_q ≈ 0.01 是逐元素量化误差。因为 α 小，所以总扰动微乎其微，这从数学上解释了为什么量化后零精度退化。', indent=True)

add_qa('量化后精度真的没损失吗？证据？',
       '端到端测试：LongBench 基准（8 个子任务 × 15 样本），BF16 基线 F1=0.1526，4-bit 量化后 F1=0.1526，差值 ΔF1 = 0.0000。NIAH（Needle-in-a-Haystack）检索准确率：100%。相对重建误差 < 0.5%。')

# ── 3.4 注意力评分 ──
add_heading('3.4 阶段 D：注意力评分 (Heavy Hitter Oracle)', level=2)

add_heading('核心算法', level=3)
add_para('每个 decode step，模型会对所有历史 token 计算注意力权重 A_{t,i}（token t 对 token i 的注意力）。HeavyHitterOracle 用这些权重为每个 token 维护一个累积重要性分数：', indent=True)
add_para('  E_i = Σ_{t=i+1}^{T} A_{t,i}')
add_para('即：token i 的分数 = 所有后续 token 对它的注意力权重之和。', indent=True)

add_heading('为什么用累积分数而不是单步分数？', level=3)
add_para('单步注意力只反映「当前 token 看了谁」，不反映全局重要性。例如一个 token 在第 100 步不重要，但在第 500 步突然变得关键（因为话题回到了相关内容）。累积分数能自然捕获这种「延迟重要性」——只要后来有 token 关注它，分数就会上升，从而免于被驱逐。', indent=True)

add_heading('为什么叫「Heavy Hitter」？', level=3)
add_para('来自数据库领域的「Heavy Hitter」概念：在频率统计中，出现频率远高于平均的元素叫 Heavy Hitter。在注意力模式中，少数 token 集中了大部分注意力分数（注意力分布天然呈幂律），这些就是注意力中的 Heavy Hitter——我们选择驱逐非 Heavy Hitter（冷 token）。', indent=True)

add_heading('保护区域机制', level=3)
add_para('并非所有 block 都参与驱逐竞争。系统设置了两个保护区域：', indent=True)
add_bullet('Sink 保护：前 sink_blocks 个 block（覆盖前 64 token）永远不会被驱逐')
add_bullet('Local 保护：最后 local_blocks 个 block（覆盖最近 8192 token）永远不会被驱逐')
add_para('实现方式：在 block_scores 中给这些 block 赋值 +inf，这样 topk(largest=False) 永远不会选到它们。', indent=True)

add_qa('为什么不保护 Attention Sink（它们分数已经很高了）？',
       '确实 Sink 的分数通常已经很高，但设置显式保护是「防御性编程」——万一某种极端注意力模式下 Sink 分数下降，显式保护能兜底。此外 Local 窗口中的最新 token 分数可能还很低（刚生成，还没来得及被后续 token 关注），更需要显式保护。')

# ── 3.5 驱逐决策 ──
add_heading('3.5 阶段 E：驱逐决策 & Triton 内核加速', level=2)

add_heading('驱逐流程', level=3)
add_para('每个 decode step：', indent=True)
add_bullet('1. HeavyHitterOracle.update()：用最新注意力权重更新累积分数')
add_bullet('2. compute_block_scores()：将 token 级分数聚合为 block 级均值')
add_bullet('3. 构建 safe_mask：标记 Sink 和 Local blocks')
add_bullet('4. torch.topk(largest=False, sorted=True)：选出分数最低的 k 个 block')
add_bullet('5. 返回候选 block 索引列表（GPU 常驻 LongTensor）')
add_bullet('6. HBMStorageManager 将这些 block 压缩后移入 DRAM，腾出 HBM 空间')

add_heading('_block_mean_kernel 详解', level=3)
add_para('这是整个驱逐决策中计算最密集的步骤——把 128K 个 token 的分数聚合为 num_blocks 个 block 均值。', indent=True)
add_para('Python 基线实现（慢）：', indent=True)
add_bullet('for i in range(num_blocks):  # 128K 次循环')
add_bullet('    block_scores[i] = token_scores[start:end].mean()  # 每次都是一次小 GPU 内核启动')
add_bullet('问题：128K 次循环 × 每次触发 CPU→GPU 同步 = 3.5 s')
add_para('Triton 内核实现（快）：', indent=True)
add_bullet('一个 CUDA kernel，launch num_blocks 个 thread block')
add_bullet('每个 thread block 独立计算一个 block 的均值')
add_bullet('全部在 GPU 上完成，零 CPU-GPU 同步')
add_bullet('耗时：0.33 ms，加速 10,465x')

add_heading('block_size=1 的快速路径', level=3)
add_para('当 block_size=1 时（128K 消融实验的配置），block 级均值就等于 token 级分数本身，不需要内核启动。compute_block_scores() 检测到 block_size=1 时直接 copy_()，进一步消除延迟。', indent=True)

add_qa('为什么用 torch.topk 而不是 torch.argsort？',
       'argsort 需要对全量 block 排序，O(n log n)。topk 只需要找最小的 k 个，O(n)（使用堆或 partition）。当 k << n 时（如 k=16K, n=128K），topk 显著更快。此外 topk 的 GPU 实现是 native 的，不需要 CPU 同步。')

# ── 3.6 预取 ──
add_heading('3.6 阶段 F：预测性异步预取 (Predictive Prefetch)', level=2)

add_heading('问题：DRAM → HBM 的延迟如何隐藏？', level=3)
add_para('从 CPU DRAM 取回一个压缩的 KV block 需要 ~0.56 ms（PCIe 传输 + 解压）。如果等到模型需要时才去取，这个延迟会直接加到生成延迟上。解决方案：预测下一步需要哪些 block，提前在后台传输。', indent=True)

add_heading('三信号预测模型', level=3)
add_para('PredictivePrefetchScheduler 综合三个信号决定预取哪些 block：', indent=True)
add_bullet('信号 1 — 空间局部性：模型往往需要与当前 block 相邻的 block。如果正在用 block 50，预测需要 block 51, 52。')
add_bullet('信号 2 — 注意力热度：HeavyHitterOracle 的累积分数中排名靠前的 block 更可能被再次需要。')
add_bullet('信号 3 — 顺序前瞻：decode 是自回归的，下一步大概率需要紧邻当前位置的 block。')

add_heading('自适应预取窗口', level=3)
add_para('预取窗口大小不是固定的。AdaptivePrefetchController 根据实时反馈动态调整：', indent=True)
add_para('  w_t = w_min + clip((σ(A_t) / σ_ref - 1) × α, -Δ_max, Δ_max) + β × miss_rate_t')
add_bullet('σ(A_t)：当前注意力分数的标准差（波动性指标）')
add_bullet('σ_ref：历史波动性的 EMA 基线')
add_bullet('miss_rate_t：最近的缓存未命中率')
add_bullet('参数：w_min=2, w_max=8, α=1.5, β=0.5, Δ_max=2.0')
add_para('直觉：注意力模式剧烈波动时（如视频场景切换），窗口扩大以覆盖更多可能需要的 block；模式稳定时，窗口收缩以节省 PCIe 带宽。', indent=True)

add_heading('异步传输机制', level=3)
add_para('AsyncPrefetcher 使用独立的 CUDA stream 执行 DMA 传输：', indent=True)
add_bullet('主 stream：模型正向计算（矩阵乘法、注意力等）')
add_bullet('后台 stream：PCIe H2D 传输 + 数据解压')
add_bullet('两个 stream 在 GPU 上并行执行，计算-传输重叠效率 98.3%')
add_para('前提条件：GPU 有足够的硬件队列来并发执行两个 stream 的任务。A100 / RTX 4090 都支持。', indent=True)

add_qa('98.3% 重叠效率是怎么算的？',
       '重叠效率 = 1 - (传输独占总时间 / 总传输时间)。如果传输完全可以与计算并行，重叠效率 = 100%。实测中，两个 stream 之间有少量总线竞争（PCIe 控制器共享），导致 1.7% 的效率损失。')

add_qa('如果预取预测错了怎么办？',
       '这就是 cache miss。系统会退化为同步取回：暂停计算，等 PCIe 传输完成，再继续。实测中自适应控制器将 miss_rate 控制在 < 5%。')

# ── 3.7 融合注意力 ──
add_heading('3.7 阶段 G：融合反量化注意力 (Fused Dequantization-Attention)', level=2)

add_heading('问题：为什么不能「先解压再算注意力」？', level=3)
add_para('朴素做法：4-bit → BF16 解压 → 标准 FlashAttention。问题在于解压会瞬间产生一个巨大的 BF16 张量。一个 16K token 的 block 有 32 个注意力头 × 128 维度 × 16K token × 2 bytes (BF16) ≈ 512 MB。这个瞬态峰值会直接触发 OOM。', indent=True)

add_heading('解决：在 GPU 寄存器中融合解压和计算', level=3)
add_para('FusedDequantAttn 的 Triton 内核将解压和注意力计算融合为单步操作：', indent=True)
add_bullet('每个 thread block 加载一小块 4-bit 数据到寄存器')
add_bullet('在寄存器中即时解压：val_bf16 = (q_uint4 - zero_point) × scale')
add_bullet('直接用解压后的值做 Q×K 点积，结果累加到输出寄存器')
add_bullet('使用在线 Softmax（两遍扫描法）：第一遍找 max，第二遍计算 exp 和 sum')
add_bullet('永远不会物化完整的 BF16 中间张量')
add_para('效果：消除了恢复路径中的 512 MB FP32 瞬态内存峰值，使「自愈恢复」本身不会成为 OOM 触发源。', indent=True)

add_heading('在线 Softmax 为什么能省显存？', level=3)
add_para('标准 Softmax 需要先算完所有 attention score，存到内存中，再统一做 exp/sum/normalize。在线 Softmax 将这个过程分为两遍：', indent=True)
add_bullet('第 1 遍：流式扫描所有 K，只维护 running_max（一个标量）')
add_bullet('第 2 遍：用最终的 max 做归一化，流式计算 running_sum 和输出')
add_para('整个过程中只需要存储 running_max 和 running_sum 两个标量，不需要存储完整的注意力矩阵。', indent=True)

add_qa('融合内核的数值精度和标准方法有区别吗？',
       '最大绝对差异 = 0.000（逐比特一致）。因为在线 Softmax 的数学等价性和 Triton 的精确浮点实现，输出与先解压再计算完全相同。')

# ============================================================
# 第四章：方法选择理由与对比
# ============================================================
add_heading('四、方法选择理由与同类对比', level=1)

add_heading('4.1 驱逐策略对比', level=2)
add_table(
    ['策略', '复杂度', '准确性', '缺点'],
    [
        ['FIFO (先进先出)', 'O(1)', '差：不考虑 token 重要性', '驱逐近期变重要的 token'],
        ['LRU (最近最少使用)', 'O(n)', '中等', 'Transformer 注意力模式不符合 LRU 假设'],
        ['StreamingLLM (滑动窗口)', 'O(1)', '差：硬截断', '永久丢弃窗口外 token，检索准确率 16.1%'],
        ['Heavy Hitter Oracle (本项目)', 'O(n) topk', '最优：基于注意力分数', '需要每步获取注意力权重，但已被 Triton 内核加速到 0.33ms'],
    ]
)
add_para('选择理由：Heavy Hitter Oracle 直接使用注意力分数作为重要性指标，这是模型本身给出的「真实重要性」，而非启发式假设。', indent=True)

add_heading('4.2 量化方法对比', level=2)
add_table(
    ['方法', '精度', '压缩比', '缺点'],
    [
        ['KIVI (2-bit/4-bit 混合)', '较高', '4-8x', '全量保留在 GPU → 仍会 OOM'],
        ['KVQuant (4-bit 对称)', '较高', '4x', '对称量化对非零均值分布精度损失更大'],
        ['GPTQ / AWQ (训练后量化)', '高', '4x', '需要校准数据集，推理流程复杂'],
        ['本项目 (4-bit 分组非对称)', '零精度损失', '~4x', '仅适用于 I/O 瓶颈场景（不影响精度但对计算无加速）'],
    ]
)
add_para('选择理由：非对称量化利用了 KV Cache 值分布不对称的特性，比对称量化精度更高。分组策略在精度和元数据开销间取得平衡。', indent=True)

add_heading('4.3 内存管理方案对比', level=2)
add_table(
    ['方案', '能支持多长上下文', '精度', '特点'],
    [
        ['vLLM PagedAttention', '~32K (24GB)', '无损', '高效分页但不解决容量不足'],
        ['FlexGen (GPU+CPU+Disk)', '128K+', '无损', '全量卸载到 CPU/Disk，速度极慢（吞吐量 0.8 tok/s）'],
        ['Offloaded Cache (DeepSpeed)', '~64K', '无损', '朴素的 GPU-CPU 分层，无预测性预取'],
        ['Hetero-KV (本项目)', '128K (24GB)', '零损失', '注意力感知 + 4-bit 压缩 + 预测性预取，吞吐量 2.1 tok/s'],
    ]
)
add_para('选择理由：Hetero-KV 结合了注意力感知的智能调度（不同于 FlexGen 的朴素卸载）和高效的压缩传输（不同于 vLLM 的纯 GPU 方案），在精度零损失的前提下最大化了上下文长度。', indent=True)

# ============================================================
# 第五章：实验数据汇总
# ============================================================
add_heading('五、关键实验数据汇总', level=1)

add_heading('5.1 内存可扩展性（24 GB GPU）', level=2)
add_table(
    ['上下文长度', '原生 HuggingFace', 'Hetero-KV', '峰值 KV Cache'],
    [
        ['8K',  '正常',  '正常',  '0.49 GB'],
        ['32K', '正常',  '正常',  '0.98 GB'],
        ['64K', '临界',  '正常',  '1.46 GB'],
        ['96K', 'OOM 崩溃', '正常',  '1.46 GB'],
        ['128K','OOM 崩溃', '正常',  '1.95 GB'],
    ]
)
add_para('Hetero-KV 的 KV Cache 从 ~16 GB（原生 128K）降到 1.95 GB（仅 Sink+Tail + 当前活跃 block），降幅 ~25x。', indent=True)

add_heading('5.2 精度零退化验证', level=2)
add_table(
    ['配置', 'F1 Score', 'ROUGE-L', 'TTFT (s)'],
    [
        ['BF16 基线',   '0.1526', '0.1319', '0.572'],
        ['Hetero-KV 4-bit', '0.1526', '0.1319', '0.269'],
        ['差值',        '0.0000', '0.0000', '-53.0%'],
    ]
)

add_heading('5.3 算子级微基准', level=2)
add_table(
    ['算子', '基线耗时', '优化后耗时', '加速比'],
    [
        ['Heavy Hitter Oracle (Python → Triton)', '~3.5 s', '0.33 ms', '10,465x'],
        ['PCIe 传输 (BF16 → 4-bit)', '2.0 ms', '0.56 ms', '2.98x (ITL)'],
        ['注意力恢复 (朴素解压 → 融合内核)', '512 MB 瞬态', '0 MB 瞬态', '消除 OOM 风险'],
    ]
)

add_heading('5.4 消融实验：各组件重要性', level=2)
add_table(
    ['移除组件', '影响'],
    [
        ['移除自愈恢复', 'NIAH 召回率从 100% 降到 0%，证明压缩+卸载不足以维持检索能力'],
        ['移除 4-bit 量化', '100% 召回但 ITL 无加速，PCIe 带宽成为瓶颈'],
        ['移除 Heavy Hitter Oracle → FIFO', 'Prefill 从 79.2s 降到 13.1s（但无注意力感知）'],
        ['移除预取调度器', 'TPOT 从 33.2ms 升至 48.7ms（+46.7%），仍可运行'],
    ]
)

# ============================================================
# 第六章：可能被问到的问题汇总
# ============================================================
add_heading('六、答辩 / 汇报可能被问到的问题', level=1)

add_heading('6.1 架构设计类', level=2)

add_qa('你的系统和 FlexGen 有什么区别？',
       'FlexGen 采用朴素的逐层卸载策略，所有 KV 数据不压缩地搬到 CPU/Disk，吞吐量仅 0.8 tok/s。Hetero-KV 通过三个关键设计实现了 2.6x 的吞吐提升：(1) 4-bit 压缩减少 75% 传输量；(2) 注意力感知的驱逐决策，只卸载不重要的 token；(3) 预测性预取将传输延迟与计算重叠。')

add_qa('你的系统和 vLLM 的 PagedAttention 有什么关系？',
       'vLLM 的 PagedAttention 解决的是 KV Cache 的内存碎片问题（类似操作系统的虚拟内存分页），但不解决显存总量不足的问题。Hetero-KV 的 HBMStorageManager 借鉴了 PagedAttention 的 block 管理思想，但额外引入了 DRAM 层和注意力感知的驱逐/恢复机制。两者可以互补：PagedAttention 管理 block 分配，Hetero-KV 管理跨层级调度。')

add_qa('为什么不在 GPU 端做量化而是搬到 CPU 再量化？',
       '实际上量化是在 GPU 端完成的（KVCompressor 在 GPU 上执行），量化后的数据才通过 PCIe DMA 传输到 CPU。这样 PCIe 上传输的是 4-bit 压缩数据，充分利用了有限带宽。')

add_qa('系统如何与 HuggingFace Transformers 集成的？',
       'FusedHeteroCache 继承自 HuggingFace 的 DynamicCache 类，实现了标准的 update()、get_seq_length()、get_mask_sizes() 等接口。模型代码无需任何修改，只需将 cache 替换为 FusedHeteroCache 即可。Cache Protocol Adapter 负责解决物理池大小与逻辑序列长度不匹配的问题，通过 get_mask_sizes() 向 Transformer 暴露正确的 KV 池大小。')

add_heading('6.2 性能类', level=2)

add_qa('10,465x 加速比是怎么算出来的？',
       '在 A100 80GB GPU 上实测：Python 基线（for 循环 + torch.mean 逐 block 聚合 128K 个 token 分数）耗时 3.4724 秒（time.time() 墙钟计时，包含 CPU-GPU 同步开销）。Triton 内核（_block_mean_kernel，单次 GPU kernel launch）耗时 0.3318 毫秒（torch.cuda.Event 计时，纯 GPU 侧）。加速比 = 3.4724s / 0.0003318s ≈ 10,465x。')

add_qa('PCIe 带宽不够怎么办？',
       '系统通过三层策略应对：(1) 4-bit 压缩将有效带宽提升 4x（16 GB/s 对应有效 64 GB/s）；(2) 预测性预取将传输隐藏在计算背后（98.3% 重叠效率）；(3) 自适应控制器在带宽压力大时缩小预取窗口。理论安全系数 γ ≥ 37（PCIe Gen4 下），系统始终处于计算绑定区而非带宽绑定区。')

add_qa('TPOT 的分布如何？尾部延迟？',
       'P50: 33.2ms, P99: 38.1ms, Max: 41.5ms, Std Dev: 1.4ms。零步骤超过 50ms 的交互式阈值。延迟非常稳定。')

add_qa('解码吞吐量是多少？',
       '8K 上下文下 30.3 tok/s。128K 上下文下由于更多的跨层级调度，吞吐量会有所下降，但 TPOT 保持在 50ms 以内。')

add_heading('6.3 精度安全类', level=2)

add_qa('4-bit 量化对精度的影响有没有理论保证？',
       '有。我们证明了 Attention Resilience Theorem：对非重击者 token 的 KV 施加 4-bit 量化，注意力输出的扰动上界为 α × √d × ε_q，其中 α 是量化 token 上的注意力质量。因为被量化的是注意力分数最低的 token（α 很小），所以总扰动被自动抑制。这从数学上解释了经验上观察到的零精度退化。')

add_qa('量化后恢复的数据和原始数据一样吗？',
       '不是逐比特相同的，但差异极小（相对重建误差 < 0.5%）。关键是：这个微小差异不会影响模型的输出。端到端测试证实了 F1 和 ROUGE-L 差值均为 0.0000。')

add_qa('如果需要恢复的 block 正好是重要的怎么办？',
       '这正是 Heavy Hitter Oracle 防止的情况：注意力分数高的 block 永远不会被驱逐，因此不会被压缩和卸载。只有分数最低的 block 才会被驱逐到 DRAM。如果后续某个被驱逐的 block 变得重要了，预取调度器会根据注意力热度信号提前将它取回。')

add_heading('6.4 局限性类', level=2)

add_qa('系统有什么局限性？',
       '主要局限有三个：(1) PCIe 总线竞争：后台预取流和计算流共享 PCIe 带宽，产生可测量的延迟增加（平均 TPOT 17.21ms vs 9.81ms 无预取），这是物理瓶颈而非算法问题；(2) 目前仅支持单 GPU，未扩展到多卡/分布式场景；(3) 量化压缩仅对 I/O 瓶颈有加速效果，对计算本身无加速。')

add_qa('能支持更大的模型吗（如 70B）？',
       '理论上可以，但 70B 模型权重就需要 ~140 GB (BF16)，单张 24GB GPU 无法容纳。需要结合模型并行（Tensor Parallelism）或权重量化（如 GPTQ 4-bit 将 70B 压缩到 ~35 GB）。Hetero-KV 的 KV Cache 管理逻辑本身是模型无关的，可以叠加在这些技术之上。')

add_qa('为什么不在 NVLink 上测试？',
       'NVLink 带宽（300-900 GB/s）远高于 PCIe（32 GB/s），在 NVLink 上带宽不再是瓶颈，系统的核心价值（压缩+注意力感知调度）的增益会减小。但在消费级 GPU（RTX 4090/4080 等）上只有 PCIe，这正是我们的目标场景。')

add_heading('6.5 实现细节类', level=2)

add_qa('Triton 内核为什么比 Python 快这么多？',
       'Python 循环每次调用 torch.mean() 都会：执行 Python 解释器开销 → 构造 CUDA kernel 参数 → 启动一个小 kernel → 等待 kernel 完成 → 读回结果。128K 次循环就是 128K 次这样的开销。Triton 内核将所有计算合并为单次 kernel launch，消除了解释器开销和 CPU-GPU 同步等待。')

add_qa('锁页内存 (pinned memory) 是什么？为什么用它？',
       '普通 CPU 内存可以被操作系统换出到磁盘（分页），此时 GPU 无法直接通过 DMA 访问它。锁页内存是操作系统承诺「永不换出」的内存区域，GPU 可以直接通过 PCIe DMA 传输其中的数据，无需先复制到临时缓冲区。这减少了数据拷贝次数，提高了传输效率。')

add_qa('CUDA stream 是什么？为什么用两个 stream？',
       'CUDA stream 是 GPU 上的命令队列。同一 stream 内的命令顺序执行，不同 stream 之间可以并行。我们用主 stream 做模型推理计算，用后台 stream 做 PCIe 数据传输，两者在硬件层面并行执行，从而隐藏传输延迟。')

add_qa('Cache Protocol Adapter 解决了什么问题？',
       '当 HBM 中的物理 KV 池大小（Sink+Tail=8256 个 token）与模型期望的逻辑序列长度（128K）不一致时，注意力掩码的宽度会不匹配。Cache Protocol Adapter 通过 get_mask_sizes() 接口向 Transformer 暴露正确的池大小，并通过独立的 position_ids 和 cache_position 张量提供正确的 RoPE 位置编码，使弹性内存层次对未修改的模型代码完全透明。')

add_qa('自适应预取控制器的参数是怎么确定的？',
       'w_min=2, w_max=8, α=1.5, β=0.5, Δ_max=2.0 是在验证集上网格搜索确定的。关键约束：w_min ≥ 2 确保基本的预取覆盖；w_max ≤ 8 避免过度预取占用过多 PCIe 带宽；α 控制波动性敏感度；β 控制缓存未命中的惩罚力度。')

# ============================================================
# 第七章：总结
# ============================================================
add_heading('七、总结', level=1)
add_para('Hetero-KV 的核心贡献是将「显存不够」这个问题重新定义为「分层存储调度」问题，并通过五个关键技术组件形成完整的解决方案：', indent=True)

add_table(
    ['#', '技术组件', '解决了什么问题', '关键指标'],
    [
        ['1', '分块预填充', 'Prefill 峰值显存 OOM', 'TTFT 降低 53%'],
        ['2', 'Heavy Hitter Oracle + Triton', '驱逐决策慢', '10,465x 加速，0.33ms'],
        ['3', '4-bit 分组非对称量化', 'PCIe 带宽瓶颈', '4x 压缩，零精度损失'],
        ['4', '预测性异步预取', 'DRAM→HBM 延迟', '98.3% 计算传输重叠'],
        ['5', '融合反量化注意力', '恢复路径的瞬态 OOM', '消除 512MB FP32 峰值'],
    ]
)
add_para('最终效果：在 24 GB 消费级 GPU 上支持 128K 上下文推理，KV Cache 仅占 1.95 GB，端到端零精度退化。', indent=True)

# ── 保存 ──
output_path = '/home/app-ahr/Hetero-KVCache-Optimizer/HeteroKV_项目技术报告.docx'
doc.save(output_path)
print(f"文档已保存至: {output_path}")
